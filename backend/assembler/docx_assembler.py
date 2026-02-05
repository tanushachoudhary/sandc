"""
Build a .docx from generated sections and apply paragraph/section formatting
extracted from a style-reference document (see project root format.py).
"""
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt


# Map applyable alignment string to python-docx enum
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _get_default_applyable() -> Dict[str, Any]:
    """Default paragraph/section spec when no style reference is provided."""
    return {
        "font_name": "Times New Roman",
        "font_size_pt": 12.0,
        "alignment": "left",
        "left_indent_in": 0.0,
        "right_indent_in": 0.0,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
        "line_spacing_rule": "single",
        "line_spacing_value": None,
        "keep_with_next": False,
        "page_break_before": False,
        "paragraph_style": "Normal",
        "section": {
            "margin_left_in": 1.0,
            "margin_right_in": 1.0,
            "margin_top_in": 1.0,
            "margin_bottom_in": 1.0,
            "page_width_in": 8.5,
            "page_height_in": 11.0,
        },
    }


def _paragraph_applyables_list(formatting_blocks: Optional[List[Dict]]) -> List[Dict]:
    """Get list of applyable specs from all paragraph-type blocks (in document order)."""
    if not formatting_blocks:
        return []
    out = []
    for block in formatting_blocks:
        if block.get("type") == "paragraph" and block.get("applyable"):
            out.append(block["applyable"])
    return out


def _first_paragraph_applyable(formatting_blocks: Optional[List[Dict]]) -> Optional[Dict]:
    """Get the applyable spec from the first paragraph-type block (for section layout / fallback)."""
    applyables = _paragraph_applyables_list(formatting_blocks)
    return applyables[0] if applyables else None


def _apply_paragraph_format(paragraph, spec: Dict[str, Any], doc: Optional[Document] = None) -> None:
    """Apply an applyable spec to a paragraph and its first run (font, indent, spacing). Set style first then direct format so direct wins."""
    # Paragraph style first (e.g. Normal, Heading 1) so direct formatting below overrides where we have explicit values
    style_name = spec.get("paragraph_style")
    if doc and style_name:
        try:
            if style_name in doc.styles:
                paragraph.style = doc.styles[style_name]
        except (KeyError, ValueError):
            pass
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    # Font
    run.font.name = spec.get("font_name") or "Times New Roman"
    run.font.size = Pt(spec.get("font_size_pt") or 12)
    # Paragraph format (indentation, spacing, alignment)
    pf = paragraph.paragraph_format
    pf.alignment = _ALIGN_MAP.get((spec.get("alignment") or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)
    pf.left_indent = Inches(spec.get("left_indent_in") or 0)
    pf.right_indent = Inches(spec.get("right_indent_in") or 0)
    pf.space_before = Pt(spec.get("space_before_pt") or 0)
    pf.space_after = Pt(spec.get("space_after_pt") or 0)
    pf.keep_with_next = spec.get("keep_with_next", False)
    pf.page_break_before = spec.get("page_break_before", False)
    # Line spacing
    rule = (spec.get("line_spacing_rule") or "single").lower()
    value = spec.get("line_spacing_value")
    if rule == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif rule == "multiple" and value is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = value
    elif rule == "exact" and value is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(value)
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _apply_section_format_to_section(section, spec: Dict[str, Any]) -> None:
    """Apply section-level formatting (margins, page size) from spec['section'] to a Section object."""
    s = spec.get("section") or {}
    section.left_margin = Inches(s.get("margin_left_in", 1.0))
    section.right_margin = Inches(s.get("margin_right_in", 1.0))
    section.top_margin = Inches(s.get("margin_top_in", 1.0))
    section.bottom_margin = Inches(s.get("margin_bottom_in", 1.0))
    section.page_width = Inches(s.get("page_width_in", 8.5))
    section.page_height = Inches(s.get("page_height_in", 11.0))


class DocxAssembler:
    """Assemble generated section text into a Word document with optional styling."""

    def __init__(self, formatting_blocks: Optional[List[Dict]] = None):
        """
        Args:
            formatting_blocks: List of block dicts from format.extract_formatting_from_file()
                               (style reference doc). Paragraph blocks should include
                               an "applyable" key. Section layout uses the first paragraph's
                               spec; each blueprint section uses the matching paragraph's
                               spec (section 0 → 1st style para, section 1 → 2nd, etc., cycling).
        """
        self._applyables_list = _paragraph_applyables_list(formatting_blocks)
        self._default_applyable = _get_default_applyable()
        # For section-level layout (margins, page size) use first applyable if any
        self._section_spec = _first_paragraph_applyable(formatting_blocks) or self._default_applyable

    def _applyable_for_section_index(self, section_index: int) -> Dict[str, Any]:
        """Return the applyable spec to use for this section (cycles if fewer style paras than sections)."""
        if not self._applyables_list:
            return self._default_applyable
        return self._applyables_list[section_index % len(self._applyables_list)]

    def _applyable_for_paragraph_index(self, paragraph_index: int) -> Dict[str, Any]:
        """Return the applyable spec for this paragraph index so 1st para matches 1st sample para, 2nd matches 2nd, etc."""
        if not self._applyables_list:
            return self._default_applyable
        return self._applyables_list[paragraph_index % len(self._applyables_list)]

    def assemble_to_docx(
        self,
        blueprint: Dict,
        sections: Dict[str, str],
        strip_leading_section_title: bool = True,
        section_formatting_map: Optional[Dict[str, int]] = None,
    ) -> bytes:
        """
        Build a Document, add section content in blueprint order with styling, return .docx bytes.

        Args:
            blueprint: Dict with "sections" list of {"name", "purpose", ...}.
            sections: Map section name -> generated text.
            strip_leading_section_title: If True, remove a leading line that matches section name.
            section_formatting_map: Optional. If provided, map section name -> applyable index (0-based).
                                   Each section uses that reference formatting. When None, uses
                                   paragraph-by-paragraph mapping (1st para -> 1st sample format, etc.).
        """
        doc = Document()
        _apply_section_format_to_section(doc.sections[0], self._section_spec)

        use_section_map = section_formatting_map and self._applyables_list
        global_para_index = 0
        for s in blueprint.get("sections", []):
            name = s.get("name", "")
            text = (sections.get(name) or "").strip()
            if strip_leading_section_title and name:
                text = self._strip_leading_section_title(text, name)
            if not text:
                continue
            if use_section_map:
                idx = section_formatting_map.get(name, 0)
                section_applyable = self._applyables_list[max(0, min(idx, len(self._applyables_list) - 1))]
            for para_text in text.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                applyable = section_applyable if use_section_map else self._applyable_for_paragraph_index(global_para_index)
                p = doc.add_paragraph()
                p.add_run(para_text)
                _apply_paragraph_format(p, applyable, doc=doc)
                global_para_index += 1

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def assemble_to_docx_stream(
        self,
        blueprint: Dict,
        sections: Dict[str, str],
        strip_leading_section_title: bool = True,
        section_formatting_map: Optional[Dict[str, int]] = None,
    ) -> BytesIO:
        """Same as assemble_to_docx but returns a BytesIO stream (e.g. for download)."""
        data = self.assemble_to_docx(
            blueprint, sections, strip_leading_section_title, section_formatting_map
        )
        stream = BytesIO(data)
        stream.seek(0)
        return stream

    @staticmethod
    def _strip_leading_section_title(text: str, section_name: str) -> str:
        """Remove a leading line if it is the section name/title."""
        if not text or not section_name:
            return text
        lines = text.strip().splitlines()
        if not lines:
            return text
        first = lines[0].strip()
        normalized_first = first.lstrip(".#0123456789) ").strip().strip("*_")
        normalized_name = section_name.strip().strip("*_")
        if normalized_first.lower() == normalized_name.lower():
            rest = "\n".join(lines[1:]).strip()
            return rest if rest else text
        return text
